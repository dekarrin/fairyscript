import scp
import copy

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Inches
from docx.shared import Pt

class DocxCompiler(object):
	def __init__(self):
		self._document = None
		self._last_paragraph = None
		self._just_completed_line = False
		self._last_speaker = None
		self._num_docs = 0
		self.screenplay_mode = False
		self.scene_metadata = None
		self.main_character = None
		self._warnings = {}
		self._screenplay_actor_margin = Inches(4)
		self.paragraph_spacing = 0
		self.include_flagsets = True
		self.include_varsets = True
		self.include_python = True
		self._add_break = True
	
	def compile_script(self, script, inputfile=None):
		self._check_format()
		self._num_docs += 1
		if inputfile is None:
			self._document = Document()
		else:
			self._document = Document(inputfile)
		self._document.add_heading("Script File #" + str(self._num_docs))
		self._last_paragraph = None
		self._last_run = None
		self._just_completed_line = False
		self._last_speaker = None
		
		if script is not None:
			for statement in script:
				self.compile_statement(statement)
		return self._document
	
	def add_warning(self, key, text):
		if key not in self._warnings:
			self._warnings[key] = []
		self._warnings[key].append(text)
		
	def get_warnings(self):
		warns = []
		for k in self._warnings:
			w_list = self._warnings[k]
			for w in w_list:
				warns.append(w)
		return warns
		
	def clear_warnings(self):
		self._warnings = {}
		
	def set_para_format(self, **kwargs):
		fmt = self._last_paragraph.paragraph_format
		for k in kwargs:
			setattr(fmt, k, kwargs[k])
			
	def set_text_format(self, **kwargs):
		fmt = self._last_run.font
		for k in kwargs:
			setattr(fmt, k, kwargs[k])
		
	def add_paragraph(self, text=None, bold=None, italic=None, style=None):
		if style is None:
			self._last_paragraph = self._document.add_paragraph()
		else:
			self._last_paragraph = self._document.add_paragraph(style=style)
		self.set_para_format(space_before=Pt(self.paragraph_spacing))
		self.set_para_format(space_after=Pt(self.paragraph_spacing))
		self._last_run = None
		if text is not None:
			self.add_run(text[0:1].upper() + text[1:], bold, italic)
			
	def add_run(self, text, bold=None, italic=None):
		self._last_run = self._last_paragraph.add_run(text)
		if bold is not None:
			self.set_text_format(bold=bold)
		if italic is not None:
			self.set_text_format(italic=italic)
		
	def compile_statement(self, statement):
		self._add_break = True
		if statement['type'] == 'line':
			self._compile_line(statement)
		else:
			if self._just_completed_line and not self.screenplay_mode:
				self.add_paragraph()
				self._last_speaker == None
				self._just_completed_line = False
			if statement['type'] == 'comment':
				self.add_paragraph('(' + scp.extract_comment(statement['text']) + ')', italic=True)
			else:
				instruction = statement['instruction']
				func = getattr(DocxCompiler, '_compile_' + instruction)
				func(self, statement)
			if self._add_break:
				self.add_paragraph()
		
	def _compile_line(self, line):
		internal = False
		sp = line['speaker']
		if sp is not None:
			sp = sp[1]
		else:
			internal = True
		continuing = self._just_completed_line and sp == self._last_speaker
		if self.screenplay_mode:
			cont = ""
			vo = ""
			charname = None
			if continuing:
				cont = " (CONT'D)"
			if internal:
				vo = " (V.O.)"
				charname = self.main_character.upper()
			else:
				charname = sp.upper()
			self.add_paragraph(charname + vo + cont)
			self.set_para_format(left_indent=self._screenplay_actor_margin)
		else:
			if self._just_completed_line and sp != self._last_speaker:
				self.add_paragraph()
			if sp is None:
				self.add_paragraph(line['text'][1])
			else:
				self.add_paragraph(sp + ': "' + line['text'][1] + '"')
		self._just_completed_line = True
		self._last_speaker = sp
		
	def _compile_SCENE(self, scene):
		if scene['transition'] is None:
			trans = "cut to:"
		else:
			trans = scene['transition'][1].lower() + " to:"
		self.add_paragraph(trans, italic=True, bold=True)
		self.set_para_format(alignment=WD_ALIGN_PARAGRAPH.RIGHT)
		self.add_paragraph()
		self.add_paragraph()
		name = scp.to_words(scene['name'][1]).title()
		self.add_paragraph(name, italic=True, bold=True)
		
	def _compile_ENTER(self, enter):
		geom = enter['motion']
		dest = None
		origin = None
		if geom is not None:
			if geom['destination'] is not None:
				dest = geom['destination'][1]
			if geom['origin'] is not None:
				origin = geom['origin'][1]
				
		line = scp.to_words(enter['target'][1]).title() + ' '
		if len(enter['states']) > 0:
			line += ','
			line += make_states(enter['states'])
			line += ', '
		if enter['transition'] is not None:
			line += scp.to_words(enter['transition'][1]).lower() + 's in to '
		else:
			line += 'enters '
		line += 'the scene'
		if origin is not None:
			line += ' near the ' + scp.to_words(origin).lower()
		if dest is not None:
			line += ' and moves to the ' + scp.to_words(dest).lower()
		if geom is not None:
			line += scp.get_duration_words(geom['duration'], 'over %d seconds')
		line += '.'
		self.add_paragraph(line, italic=True)

	def _compile_ACTION(self, action):
		line = scp.to_words(action['target'][1]).title() + ' '
		if len(action['states']) > 0:
			line += 'changes to appear'
			line += make_states(action['states'])
			if action['destination'] is not None:
				line += ', and then '
		if action['destination'] is not None:
			line += 'moves to the ' + scp.to_words(action['destination'][1]).lower()
			line += scp.get_duration_words(action['duration'], 'over %d seconds')
		line += '.'
		self.add_paragraph(line, italic=True)
		
	def _compile_EXIT(self, exit):
		geom = exit['motion']
		dest = None
		origin = None
		if geom is not None:
			if geom['destination'] is not None:
				dest = geom['destination'][1]
			if geom['origin'] is not None:
				origin = geom['origin'][1]
		
		line = scp.to_words(exit['target'][1]).title() + ' '
		if geom is not None:
			line += 'moves '
		if origin is not None:
			line += 'from the ' + scp.to_words(origin).lower() + ' '
		if dest is not None:
			line += 'to the ' + scp.to_words(dest).lower() + ' '
		if geom is not None:
			line += 'and '
		if exit['transition'] is not None:
			line += scp.to_words(exit['transition'][1]).lower() + 's out of '
		else:
			line += 'exits '
		line += 'the scene'
		if geom is not None:
			line += scp.get_duration_words(geom['duration'], 'over %d seconds')
		self.add_paragraph(line, italic=True)
		
	def _compile_MUSIC(self, music):
		self.add_paragraph()
		line = 'We hear '
		if music['action'] == 'start':
			if music['fadeout'] is not None:
				line += 'the current music fade out' + scp.get_duration_words(music['fadeout'], 'over %d seconds') + ', '
				line += 'and then we hear '
			line += 'the song '
			self.add_run(line, italic=True)
			self.add_run(scp.to_words(music['target'][1]).title(), italic=False)
			self.add_run(' begin to play.', italic=True)
		elif music['action'] == 'stop':
			if scp.typed_check(music['target'], 'rel'):
				line += music['target'][1].lower() + ' music '
			else:
				line += 'the song '
				self.add_run(line, italic=True)
				self.add_run(scp.to_words(music['target'][1]).title(), italic=False)
				line = ' '
			if music['duration'] is not None:
				line += 'fade out' + scp.get_duration_words(music['duration'], 'over %d seconds') + '.'
			else:
				line += 'stop.'
			self.add_run(line, italic=True)
			
	def _compile_GFX(self, gfx):
		line = 'We see '
		if gfx['action'] == 'start':
			a = scp.indef_article(gfx['target'][1])
			line += a + ' ' + scp.to_words(gfx['target'][1]).upper()
			if scp.typed_check(gfx['loop'], 'boolean', True):
				line += ' effect begin and continue'
			line += '.'
		elif gfx['action'] == 'stop':
			if scp.typed_check(gfx['target'], 'rel'):
				line += gfx['target'][1].lower() + ' effects '
			else:
				line += 'the ' + scp.to_words(gfx['target'][1]).upper() + ' effect '
			if gfx['duration'] is not None:
				line += 'fade away' + scp.get_duration_words(gfx['duration'], 'over %d seconds') + '.'
			else:
				line += 'stop.'
		self.add_paragraph(line, italic=True)

	def _compile_SFX(self, sfx):
		line = 'We hear '
		if sfx['action'] == 'start':
			a = scp.indef_article(sfx['target'][1])
			line += a + ' ' + scp.to_words(sfx['target'][1]).upper()
			if scp.typed_check(sfx['loop'], 'boolean', True):
				line += ' sound begin to repeat'
			line += '.'
		elif sfx['action'] == 'stop':
			if scp.typed_check(sfx['target'], 'rel'):
				line += sfx['target'][1].lower() + ' repeating sounds '
			else:
				line += 'the repeated ' + scp.to_words(sfx['target'][1]).upper() + ' sound '
			if sfx['duration'] is not None:
				line += 'fade away' + scp.get_duration_words(sfx['duration'], 'over %d seconds') + '.'
			else:
				line += 'stop.'
		self.add_paragraph(line, italic=True)
			
	def _compile_FMV(self, fmv):
		self.add_paragraph()
		self.add_run("We see the full-motion video ", italic=True)
		self.add_run(scp.to_words(fmv['target'][1]), italic=False)
		self.add_run(' play.')
		
	def _compile_CAMERA(self, camera):
		line = 'We'
		acts_added = 0
		for act in camera['actions']:
			if acts_added == len(camera['actions']) - 1 and len(camera['actions']) > 1:
				line += ' and then'
			if 'duration' in act and scp.typed_check(act['duration'], 'rel'):
				line += scp.get_duration_words(act['duration'], '')
			if act['type'] == 'SNAP':
				line += ' focus on the ' + scp.to_words(act['target'][1])
			elif act['type'] == 'PAN':
				line += ' shift our focus to the ' + scp.to_words(act['target'][1])
			elif act['type'] == 'ZOOM':
				line += ' move '
				if scp.typed_check(act['target'], 'rel', 'IN'):
					line += ' in closer'
				elif scp.typed_check(act['target'], 'rel', 'OUT'):
					line += ' back farther'
			if 'duration' in act and not scp.typed_check(act['duration'], 'rel'):
				line += scp.get_duration_words(act['duration'], 'over %d seconds')
			if acts_added < len(camera['actions']) - 1 and len(camera['actions']) > 2:
				line += ','
			acts_added += 1
		self.add_paragraph(line, italic=True)
		
	def _compile_CHOICE(self, choice):
		if choice['label'] is not None:
			labelstmt = {'type': 'annotation', 'instruction': 'SECTION', 'section': choice['label'], 'params': []}
			self.compile_statement(labelstmt)
		line = "We are presented with a choice:"
		if choice['title'] is not None:
			line += ' ' + choice['title'][1]
		self.add_paragraph(line)
		for c in choice['choices']:
			self.add_paragraph(style='List Bullet')
			if c['condition'] is not None:
				if scp.typed_check(c['condition'], 'boolean', True):
					self.add_run('(always shown) ', italic=True)
				elif scp.typed_check(c['condition'], 'boolean', False):
					self.add_run('(never shown) ', italic=True)
				elif scp.typed_check(c['condition'], 'id'):
					cond_id = scp.to_words(c['condition'][1]).lower()
					if cond_id.startswith('have '):
						self.add_run('(only shown if we ' + cond_id + ') ', italic=True)
					else:
						self.add_run('(only shown if \'' + cond_id + '\' is set) ', italic=True)
				else:
					self.add_run('(only shown if ' + scp.to_human_readable(c['condition'][1]) + ') ', italic=True)
				line = c['text'][1] + ":\n"
				for v in c['sets']:
					line += self.make_varset(v) + "\n"
				line += self.make_goto({'destination': c['target']})
				self.add_run(line, italic=False)
			
	def _compile_DESCRIPTION(self, desc):
		line = ''
		if desc['target'] is not None:
			line = "Regarding " + scp.to_words(desc['target'][1]) + ': '
		line += desc['text'][1]
		self.add_paragraph(line)
			
	def _compile_SECTION(self, section):
		self._document.add_heading(scp.to_words(section['section'][1]).title(), level=2)
		
	def _compile_FLAGSET(self, flagset):
		if self.include_flagsets:
			self.add_paragraph(self.make_flagset(flagset), italic=True)
		else:
			self._add_break = False
				
	def _compile_VARSET(self, varset):
		if self.include_varsets:
			self.add_paragraph(self.make_varset(varset), italic=True)
		else:
			self._add_break = False
		
	def _compile_DIALOG(self, dialog):
		self.add_paragraph("We set the dialog window to " + dialog['mode'].upper() + " mode.", italic=True)
		
	def _compile_GOTO(self, goto):
		self.add_paragraph(self.make_goto(goto))
		
	def _compile_EXECUTE(self, execute):
		line = "We set the proper parameters and execute section " + scp.to_words(execute['section'][1]).title() + '.'
		self.add_paragraph(line, italic=True)
		
	def _compile_END(self, end):
		if 'retval' in end:
			self.add_paragraph('We return the appropriate value from this section.', italic=True, bold=True)
		else:
			self._add_break = False
		
	def _compile_WHILE(self, whilestmt):
		line = "We do the following "
		cond = scp.get_expr(whilestmt['condition'])
		if scp.typed_check(whilestmt['condition'], 'boolean'):
			if cond:
				line += 'forever:'
			else:
				line = 'We never do the following:'
		elif scp.typed_check(whilestmt['condition'], 'id'):
			cond_words = scp.to_words(whilestmt['condition'][1]).lower()
			if cond_words.startswith('have '):
				line += 'while we ' + cond_words + ':'
			else:
				line += 'while ' + scp.quote(cond_words, "'") + ' is set:'
		else:
			line += 'while ' + scp.to_human_readable(cond) + ':'
		self.add_paragraph(line, italic=True)
		self.add_paragraph('{', italic=True)
		self.add_paragraph()
		for st in whilestmt['statements']:
			self.compile_statement(st)
		self.add_paragraph('}', italic=True)
		
	def _compile_IF(self, ifstmt):
		elsebr = None
		firstbr = True
		for br in ifstmt['branches']:
			if br['condition'] is None:
				elsestmt = br
			else:
				negation = ''
				if firstbr:
					line = 'We%s do the following' + scp.get_expr(br['condition']) + ':'
					firstbr = False
				else:
					line = 'Otherwise, we%s do the following'
				cond = scp.get_expr(br['condition'])
				if scp.typed_check(br['condition'], 'boolean'):
					if cond:
						negation = ' always'
						line += ':'
					else:
						negation = ' never'
						line += ':'
				elif scp.typed_check(br['condition'], 'id'):
					cond_words = scp.to_words(br['condition'][1]).lower()
					if cond_words.startswith('have '):
						line += ' if we ' + cond_words + ':'
					else:
						line += ' if ' + scp.quote(cond_words, "'") + ' is set:'
				else:
					line += ' if ' + scp.to_human_readable(cond) + ':'
				self.add_paragraph(line % negation, italic=True)
				self.add_paragraph('{', italic=True)
				self.add_paragraph()
				for st in br['statements']:
					self.compile_statement(st)
				self.add_paragraph('}', italic=True)
				self.add_paragraph()
		if elsebr is not None:
			self.add_paragraph("Otherwise, we do the following:", italic=True)
			self.add_paragraph("{", italic=True)
			self.add_paragraph()
			for st in elsebr['statements']:
				self.compile_statement(st)
			self.add_paragraph("}", italic=True)
			self.add_paragraph()
		self._add_break = False
			
	def _compile_PYTHON(self, python):
		if self.include_python:
			self.add_paragraph("We execute the following python code:", italic=True)
			self.add_paragraph("{", italic=True)
			self.add_paragraph()
			lines = python['body'].split('\n')
			for line in lines:
				self.add_paragraph(line.strip())
			self.add_paragraph("}", italic=True)
		else:
			self.add_paragraph("We execute python code.", italic=True)
		
	def _check_format(self):
		if self.screenplay_mode:
			if self.main_character is None:
				self.add_warning('main_char_not_defined', "a main character is required for screenplay mode; using 'Main character'")
				self.main_character = 'Main character'
				
	def make_flagset(self, flagset):
		line = ''
		value = scp.get_expr(flagset['value'])
		flag = scp.to_words(flagset['name'][1]).lower()
		if scp.typed_check(flagset['value'], 'boolean'):
			if not flag.startswith('have '):
				if value:
					line = 'We set ' + scp.quote(flag, "'") + '.'
				else:
					line = 'We unset ' + scp.quote(flag, "'") + '.'
			elif value:
				line = 'We now ' + flag + '.'
			else:
				line = 'We now no longer ' + flag + '.'
		elif scp.typed_check(flagset['value'], 'id'):
			if flag.startswith('have '):
				line = 'Whether we ' + flag + ' is determined by '
			else:
				line = 'We set ' + scp.quote(flag, "'") + ' to the same as '
			value_id = scp.to_words(value).lower()
			if value_id.startswith('have '):
				line += 'whether we ' + value_id + '.'
			else:
				line += 'the value of ' + scp.quote(value_id, "'") + '.'
		elif scp.typed_check(flagset['value'], 'expr'):
			if flag.startswith('have '):
				line = 'Whether we now ' + flag + ' is determined by the value of ' + value + '.'
			else:
				line = 'We set ' + scp.quote(flag, "'") + ' to the same as the value of ' + value + '.'
		return line
		
	def make_varset(self, varset):
		line = ''
		if scp.typed_check(varset['value'], 'boolean'):
			line = self.make_flagset(varset)
		var = scp.to_words(varset['name'][1]).lower()
		value = scp.get_expr(varset['value'])
		if scp.typed_check(varset['value'], 'incdec'):
			line = 'We ' + scp.to_human_readable(var + value).replace(var, scp.quote(var, "'")) + '.'
		else:
			line = 'We set ' + scp.quote(var, "'") + ' to ' + value + '.'
		return line
		
	def make_goto(self, goto):
		return 'We jump to section ' + scp.to_words(goto['destination'][1]).title() + '.'
				
def make_states(states):	
	states_added = 0
	line = ''
	for s in states:
		if states_added == len(states) - 1 and len(states) > 1: # we are on the last one
			line += ' and'
		line += ' ' + scp.to_words(s[1]).lower()
		if states_added < len(states) - 1 and len(states) > 2:
			line += ','
		states_added += 1
	return line